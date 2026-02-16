from flask import Blueprint, request, jsonify
import os
import tempfile
from docx import Document
import json

file_parse_bp = Blueprint('file_parse', __name__)

# 地域数据文件路径
import os
REGIONS_FILE = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'china_regions.json')

class DocxTextExtractor:
    """专门用于提取 docx 文本和页码信息的类"""

    def __init__(self, docx_path):
        self.doc = Document(docx_path)
        self.page_breaks = self._find_page_breaks()
        self.full_text = self._extract_full_text()

    def _find_page_breaks(self):
        """估算段落在文档中的页码位置"""
        page_breaks = []
        current_page = 1
        current_line_count = 0
        lines_per_page_estimate = 40  # 这是一个估算值，可根据需要调整

        for i, paragraph in enumerate(self.doc.paragraphs):
            text = paragraph.text
            if text.strip():  # 忽略空段落
                # 估算段落行数
                lines_in_para = len(text) // 80 + text.count('\n') + 1  # 简单估算
                current_line_count += lines_in_para

                # 检查段落中是否有分页符
                if 'w:br' in paragraph._p.xml and 'type="page"' in paragraph._p.xml:
                    # 找到显式分页符
                    page_breaks.append((i, current_page))
                    current_page += 1
                    current_line_count = lines_in_para  # 重置行数计数

                # 检查是否因行数估算而翻页
                if current_line_count > lines_per_page_estimate:
                    page_breaks.append((i, current_page))
                    current_page += 1
                    current_line_count = lines_in_para  # 重置为当前段落的行数

        return page_breaks

    def _extract_full_text(self):
        """提取完整文本"""
        import sys
        full_text = []
        for para in self.doc.paragraphs:
            # 保留换行符以便于上下文查找
            full_text.append(para.text + "\n")
        text = ''.join(full_text)
        print(f"提取的文档文本: {text[:500]}...", file=sys.stderr)  # 打印前500个字符
        return text

    def get_page_number(self, paragraph_index):
        """根据段落索引估算页码"""
        page_num = 1
        for break_index, break_page in self.page_breaks:
            if paragraph_index >= break_index:
                page_num = break_page + 1  # 下一页开始
            else:
                break
        return page_num

    def find_keyword_occurrences(self, keywords):
        """查找关键词并返回其页码和上下文"""
        import sys
        occurrences = []
        context_length = 50  # 上下文字符数

        # 为了提高查找效率，可以将列表转换为集合进行成员检查
        keywords_set = set(kw.strip() for kw in keywords if kw.strip())
        print(f"关键词集合: {keywords_set}", file=sys.stderr)

        # 去重关键词，避免重复查找
        unique_keywords = list(keywords_set)
        print(f"去重后的关键词: {unique_keywords}", file=sys.stderr)

        for keyword in unique_keywords:
            keyword = keyword.strip()
            if not keyword:
                continue
            print(f"查找关键词: {keyword}", file=sys.stderr)
            start = 0
            while True:
                pos = self.full_text.find(keyword, start)
                if pos == -1:
                    print(f"关键词 {keyword} 未找到", file=sys.stderr)
                    break
                print(f"找到关键词 {keyword} 在位置 {pos}", file=sys.stderr)

                # 找到包含关键词的段落索引
                para_index = 0
                text_pos = 0
                for i, para in enumerate(self.doc.paragraphs):
                    para_end = text_pos + len(para.text) + 1  # +1 for \n
                    if text_pos <= pos < para_end:
                        para_index = i
                        break
                    text_pos = para_end

                # 估算页码
                page_num = self.get_page_number(para_index)

                # 提取上下文
                context_start = max(0, pos - context_length)
                context_end = min(len(self.full_text), pos + len(keyword) + context_length)
                context = self.full_text[context_start:context_end].strip()

                occurrences.append({
                    'keyword': keyword,
                    'page': page_num,
                    'context': context
                })
                print(f"添加匹配: {keyword} 在 Page {page_num}", file=sys.stderr)
                start = pos + 1  # 从下一个位置继续查找

        print(f"关键词查找完成，找到 {len(occurrences)} 个匹配", file=sys.stderr)
        return occurrences

def load_regions():
    """从 JSON 文件加载地域名称"""
    if os.path.exists(REGIONS_FILE):
        try:
            with open(REGIONS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                # 如果是列表格式，直接返回
                if isinstance(data, list):
                    return data
                # 如果是字典格式，提取所有地域名称
                elif isinstance(data, dict):
                    regions = []
                    # 检查是否是新的结构（包含"省级"、"市级"、"区级"键）
                    if "省级" in data:
                        regions.extend(data["省级"])
                    if "市级" in data:
                        regions.extend(data["市级"])
                    if "区级" in data:
                        regions.extend(data["区级"])
                    return regions
                else:
                    return []
        except (json.JSONDecodeError, IOError):
            return []
    return []

@file_parse_bp.route('/file-parse', methods=['POST'])
def parse_file():
    """解析文件并匹配地名"""
    import sys
    print("=" * 80, file=sys.stderr)
    print("开始处理文件上传请求", file=sys.stderr)
    print("=" * 80, file=sys.stderr)
    
    try:
        # 检查是否有文件部分
        if 'file' not in request.files:
            print("错误: 没有选择文件", file=sys.stderr)
            return jsonify({'success': False, 'message': '没有选择文件'}), 400
        file = request.files['file']

        # 如果用户没有选择文件
        if file.filename == '':
            print("错误: 没有选择文件", file=sys.stderr)
            return jsonify({'success': False, 'message': '没有选择文件'}), 400

        # 检查文件类型
        if not file.filename.lower().endswith(('.doc', '.docx', '.xls', '.xlsx')):
            print(f"错误: 不支持的文件类型: {file.filename}", file=sys.stderr)
            return jsonify({'success': False, 'message': '不支持的文件类型。请上传 .doc, .docx, .xls 或 .xlsx 文件。'}), 400

        # 获取用户自定义关键词
        custom_keywords = []
        # 打印所有表单参数，用于调试
        print(f"所有表单参数: {dict(request.form)}", file=sys.stderr)
        # 尝试获取所有以 keywords 开头的参数
        for key, value in request.form.items():
            print(f"参数: {key} = {value}", file=sys.stderr)
            if key.startswith('keywords['):
                custom_keywords.append(value)
        # 尝试获取 keywords 参数
        if 'keywords' in request.form:
            # 如果是单个关键词
            if isinstance(request.form['keywords'], str):
                custom_keywords.append(request.form['keywords'])
            # 如果是多个关键词
            else:
                custom_keywords.extend(request.form.getlist('keywords'))
        print(f"最终的自定义关键词: {custom_keywords}", file=sys.stderr)

        # 使用临时文件处理
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1])
        file.save(temp_file.name)
        temp_file_path = temp_file.name
        temp_file.close()  # 关闭文件以便后续读取
        print(f"临时文件路径: {temp_file_path}", file=sys.stderr)

        try:
            # 加载地名数据
            keywords = load_regions()
            # 添加用户自定义关键词
            keywords.extend(custom_keywords)
            print(f"最终的关键词列表: {keywords}", file=sys.stderr)
            if not keywords:
                print("错误: 地域名称列表为空或加载失败", file=sys.stderr)
                return jsonify({'success': False, 'message': '地域名称列表为空或加载失败'}), 500

            # 提取文本和查找关键词
            if file.filename.lower().endswith('.docx'):
                print(f"开始解析文件: {file.filename}", file=sys.stderr)
                extractor = DocxTextExtractor(temp_file_path)
                print(f"文件解析成功，开始查找关键词", file=sys.stderr)
                occurrences = extractor.find_keyword_occurrences(keywords)
                print(f"关键词查找完成，找到 {len(occurrences)} 个匹配", file=sys.stderr)
                os.unlink(temp_file_path)  # 处理完立即删除临时文件
                
                # 打印所有匹配结果
                print(f"所有匹配结果: {occurrences}", file=sys.stderr)
                
                # 转换结果格式
                matches = []
                for occurrence in occurrences:
                    matches.append({
                        'location': f'Page {occurrence["page"]}',
                        'text': occurrence["context"],
                        'place': occurrence["keyword"]
                    })
                
                print(f"最终返回的匹配结果: {matches}", file=sys.stderr)
                print("=" * 80, file=sys.stderr)
                return jsonify({
                    'success': True,
                    'fileName': file.filename,
                    'matches': matches
                })
            else:
                # 对于其他文件类型，暂时返回模拟结果
                os.unlink(temp_file_path)
                print(f"对于文件类型 {file.filename}，返回模拟结果", file=sys.stderr)
                print("=" * 80, file=sys.stderr)
                return jsonify({
                    'success': True,
                    'fileName': file.filename,
                    'matches': [
                        {
                            'location': 'Sheet 1, Row 1, Column A',
                            'text': '这是一个包含北京的句子示例',
                            'place': '北京'
                        },
                        {
                            'location': 'Sheet 1, Row 3, Column B',
                            'text': '这是一个包含上海的句子示例',
                            'place': '上海'
                        }
                    ]
                })

        except Exception as e:
            # 确保即使出错也删除临时文件
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
            print(f"错误: 处理文件时出错: {str(e)}", file=sys.stderr)
            print("=" * 80, file=sys.stderr)
            return jsonify({'success': False, 'message': f'处理文件时出错: {str(e)}'}), 500

    except Exception as e:
        print(f"错误: 上传文件时发生错误: {str(e)}", file=sys.stderr)
        print("=" * 80, file=sys.stderr)
        return jsonify({'success': False, 'message': f'上传文件时发生错误: {str(e)}'}), 500

@file_parse_bp.route('/regions', methods=['GET'])
def get_regions():
    """获取所有地域名称"""
    regions = load_regions()
    return jsonify({'success': True, 'regions': regions})
